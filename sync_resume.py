import json
import re
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Paths
WORKSPACE_DIR = os.path.dirname(os.path.abspath(__file__))
JSON_PATH = os.path.join(WORKSPACE_DIR, "resume_data.json")
HTML_PATH = os.path.join(WORKSPACE_DIR, "index.html")
DOCX_PATH = os.path.join(WORKSPACE_DIR, "assets", "Sanket_Repale_Resume.docx")

def replace_block(content, start_tag, end_tag, replacement):
    pattern = re.compile(rf"{start_tag}.*?{end_tag}", re.DOTALL)
    return pattern.sub(f"{start_tag}\n{replacement}\n{end_tag}", content)

def update_html(data):
    print("Reading index.html...")
    with open(HTML_PATH, "r", encoding="utf-8") as f:
        html = f.read()

    # Bio
    bio_html = data["summary"]
    html = replace_block(html, "<!-- START_BIO -->", "<!-- END_BIO -->", f"                    {bio_html}")

    # ZS Experience
    zs_exp = data["experience"][0]
    zs_bullets_html = '\n'.join([f'                                            <li>{b}</li>' for b in zs_exp["bullets"]])
    zs_html = f'''                                        <ul class="exp-list" style="list-style-type: disc;">
{zs_bullets_html}
                                        </ul>'''
    html = replace_block(html, "<!-- START_ZS_EXPERIENCE -->", "<!-- END_ZS_EXPERIENCE -->", zs_html)

    # LTI Experience
    lti_exp = data["experience"][1]
    lti_bullets_html = '\n'.join([f'                                            <li>{b}</li>' for b in lti_exp["bullets"]])
    lti_html = f'''                                        <ul class="exp-list" style="list-style-type: disc;">
{lti_bullets_html}
                                        </ul>'''
    html = replace_block(html, "<!-- START_LTI_EXPERIENCE -->", "<!-- END_LTI_EXPERIENCE -->", lti_html)

    # Certifications
    cert_list_html = '\n'.join([f'                                    <li>{c}</li>' for c in data["certifications"]])
    cert_html = f'''                                <ul class="exp-list" style="list-style-type: disc; font-size: 0.9rem;">
{cert_list_html}
                                </ul>'''
    html = replace_block(html, "<!-- START_CERTIFICATIONS -->", "<!-- END_CERTIFICATIONS -->", cert_html)

    # Awards
    award_list_html = '\n'.join([f'                                    <li>{a}</li>' for a in data["awards"]])
    award_html = f'''                                <ul class="exp-list" style="list-style-type: disc; font-size: 0.9rem;">
{award_list_html}
                                </ul>'''
    html = replace_block(html, "<!-- START_AWARDS -->", "<!-- END_AWARDS -->", award_html)

    print("Writing index.html...")
    with open(HTML_PATH, "w", encoding="utf-8") as f:
        f.write(html)
    print("index.html updated successfully.")

def add_heading_with_bottom_border(doc, text, color_hex="005B82"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    
    # Add bottom border in XML
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')  # size of border
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), color_hex)
    pbdr.append(bottom)
    pPr.append(pbdr)
    
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string(color_hex)
    return p

def parse_and_add_text_with_formatting(p, text, size=9.5):
    # Regex to split on <strong> and </strong>
    parts = re.split(r'(<strong>.*?</strong>)', text)
    for part in parts:
        if part.startswith('<strong>') and part.endswith('</strong>'):
            clean_part = part[8:-9]
            r = p.add_run(clean_part)
            r.bold = True
        else:
            r = p.add_run(part)
        r.font.size = Pt(size)
        r.font.name = 'Calibri'
        r.font.color.rgb = RGBColor.from_string("333333")

def generate_docx(data):
    print("Generating Word Document Resume...")
    doc = Document()
    
    # Set standard margins (0.75 in)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        
    # Title & Contact
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(2)
    
    run_name = title_p.add_run(data["name"].upper() + "\n")
    run_name.bold = True
    run_name.font.name = 'Calibri'
    run_name.font.size = Pt(18)
    run_name.font.color.rgb = RGBColor.from_string("005B82")
    
    run_title = title_p.add_run(data["title"])
    run_title.italic = True
    run_title.font.name = 'Calibri'
    run_title.font.size = Pt(11)
    run_title.font.color.rgb = RGBColor.from_string("555555")
    
    contact_p = doc.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_p.paragraph_format.space_after = Pt(10)
    
    contact_info = f"{data['location']}  |  {data['phone']}  |  {data['email']}\nLinkedIn: {data['linkedin']}"
    run_contact = contact_p.add_run(contact_info)
    run_contact.font.name = 'Calibri'
    run_contact.font.size = Pt(9.5)
    run_contact.font.color.rgb = RGBColor.from_string("333333")
    
    # Summary
    add_heading_with_bottom_border(doc, "PROFESSIONAL SUMMARY")
    summary_p = doc.add_paragraph()
    summary_p.paragraph_format.space_after = Pt(6)
    summary_p.paragraph_format.line_spacing = 1.15
    run_summary = summary_p.add_run(data["summary"])
    run_summary.font.name = 'Calibri'
    run_summary.font.size = Pt(9.5)
    run_summary.font.color.rgb = RGBColor.from_string("333333")
    
    # Skills
    add_heading_with_bottom_border(doc, "TECHNICAL SKILLS")
    skills_data = [
        ("Programming Languages", "C#, Java, Python"),
        ("Test Automation", "Selenium, Playwright, SpecFlow BDD, Tosca DI (Data Integrity)"),
        ("API Testing", "Ready API, Postman, Locust (Performance), RESTful and SOAP API Testing"),
        ("DevOps & CI/CD", "Azure DevOps, CI/CD Pipeline Integration, Git, GitHub"),
        ("Software Testing", "Functional, End-to-End, Regression, ETL Testing"),
        ("Development Tools", "Visual Studio, Eclipse, VS Code")
    ]
    for category, details in skills_data:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Inches(0.15)
        
        r1 = p.add_run(f"•  {category}: ")
        r1.bold = True
        r1.font.name = 'Calibri'
        r1.font.size = Pt(9.5)
        r1.font.color.rgb = RGBColor.from_string("333333")
        
        r2 = p.add_run(details)
        r2.font.name = 'Calibri'
        r2.font.size = Pt(9.5)
        r2.font.color.rgb = RGBColor.from_string("333333")
        
    # Experience
    add_heading_with_bottom_border(doc, "PROFESSIONAL EXPERIENCE")
    for job in data["experience"]:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.keep_with_next = True
        
        # Layout: Company | Role  [Right Aligned: Date]
        r_comp = p.add_run(f"{job['company']}  |  {job['role']}")
        r_comp.bold = True
        r_comp.font.name = 'Calibri'
        r_comp.font.size = Pt(10)
        r_comp.font.color.rgb = RGBColor.from_string("005B82")
        
        # Tab Stop at 7.0 inches for right alignment of dates
        p.paragraph_format.tab_stops.add_tab_stop(Inches(7.0), alignment=2)
        r_date = p.add_run(f"\t{job['duration']}")
        r_date.bold = True
        r_date.font.name = 'Calibri'
        r_date.font.size = Pt(9.5)
        r_date.font.color.rgb = RGBColor.from_string("555555")
        
        for bullet in job["bullets"]:
            p_bullet = doc.add_paragraph(style='List Bullet')
            p_bullet.paragraph_format.space_after = Pt(2.5)
            p_bullet.paragraph_format.left_indent = Inches(0.25)
            parse_and_add_text_with_formatting(p_bullet, bullet, size=9.5)
            
    # Certifications
    add_heading_with_bottom_border(doc, "CERTIFICATIONS")
    for cert in data["certifications"]:
        p_cert = doc.add_paragraph(style='List Bullet')
        p_cert.paragraph_format.space_after = Pt(2.5)
        p_cert.paragraph_format.left_indent = Inches(0.25)
        parse_and_add_text_with_formatting(p_cert, cert, size=9.5)
        
    # Awards
    add_heading_with_bottom_border(doc, "AWARDS & RECOGNITION")
    for award in data["awards"]:
        p_award = doc.add_paragraph(style='List Bullet')
        p_award.paragraph_format.space_after = Pt(2.5)
        p_award.paragraph_format.left_indent = Inches(0.25)
        parse_and_add_text_with_formatting(p_award, award, size=9.5)

    print(f"Saving DOCX resume to {DOCX_PATH}...")
    doc.save(DOCX_PATH)
    print("DOCX generated successfully.")

if __name__ == "__main__":
    main_data = None
    with open(JSON_PATH, "r", encoding="utf-8") as f:
        main_data = json.load(f)
    update_html(main_data)
    generate_docx(main_data)
    print("\nSUCCESS: Website and Downloadable Word Resume synchronized!")
